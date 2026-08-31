# app.py
# Phase 6.3 - the HTTP layer. Thin on purpose: every hard decision already lives in
# agent.py, guards.py, auth.py, db.py and telemetry.py, and this file's job is to carry
# requests to them without inventing anything of its own.
#
# THE CONTRACT THIS FILE EXISTS TO KEEP (written in PROJECT_TRACKER before any code):
# **the API returns the same answer object the eval scores.** /ask calls the exact function
# run_eval.py calls - agent.run_agent - inside the exact capture run_eval.py uses. If the
# two could drift, then what was measured is not what shipped, and every number in that
# tracker would be describing a system nobody runs.
#
# 🔴 NO STREAMING, AND THAT IS A DECISION, NOT AN OMISSION.
# guards.scrub() needs the WHOLE answer: leak detection is an 8-gram test over the finished
# text, and the token strip rewrites it. You cannot stream an answer you have not finished
# checking - the first token is out of the door before the check exists. Streaming is a feel;
# a leaked instruction is a defect. So: buffer, check, then release. The UI compensates with
# a live pipeline trace, which is more honest and more interesting than a typewriter effect.
#
# Run:  uvicorn app:app --reload --port 8000
# Test: python test_app.py     (no API key, no network, no cost)

import json
import os
import secrets
import threading
from typing import List, Optional

from dotenv import load_dotenv
from fastapi import Cookie, Depends, FastAPI, Header, HTTPException, Request, Response
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from html import escape as _x
from pydantic import BaseModel, Field

import auth
import cache
import db
import version

load_dotenv()

SESSION_COOKIE = "fa_session"
CSRF_COOKIE = "fa_csrf"
CSRF_HEADER = "x-csrf-token"

# Secure=True means the browser only sends the cookie over HTTPS, which would make local
# development over http://localhost silently log you out on every request - a failure that
# looks like a bug in the auth code. So it is configurable and it DEFAULTS TO SECURE; the
# developer opts down, and Phase 7 cannot forget to opt up.
COOKIE_SECURE = os.environ.get("COOKIE_INSECURE", "").lower() not in ("1", "true", "yes")

app = FastAPI(title="Financial Research & Compliance Analyst", version="6.3")


# --- the agent, imported once ---------------------------------------------------------------
# Imported at module level so the index is read and FILINGS is populated before the first
# request, rather than making one unlucky user pay for start-up.
#
# telemetry.install() must happen BEFORE any answer is produced, and it must reach every
# module that did `from judges import log_cost`. The assertion is the same one run_eval.py
# makes, for the same reason: a capture that reaches three modules of five under-reports
# cost and looks exactly like a cheap system.
import agent            # noqa: E402  (after load_dotenv, which it needs)
import judges           # noqa: E402
import rag              # noqa: E402
import telemetry        # noqa: E402

import rewriter      # noqa: E402  - holds its own log_cost binding; see run_eval.py

_TRACKED = (judges, rag, agent, rewriter)
telemetry.install(*_TRACKED)
# Assert the STATE, not the change. `assert len(install(...)) >= 3` was here and it was wrong:
# install() returns what this call patched, so a second import of this module - which is
# exactly what `python app.py` causes - patched nothing, and the assertion read that correct
# no-op as "cost tracking reached only []" and refused to start. unpatched() answers the
# question that was always meant: is every module that can make a paid call routed through
# the capture right now, whoever installed it.
assert not telemetry.unpatched(*_TRACKED), \
    f"cost tracking never reached {telemetry.unpatched(*_TRACKED)} - spend would be under-reported"

agent.VERBOSE = False

# The API must never serve with the guardrails off. `--no-guards` exists so the eval can
# reproduce the pre-5.1 system; there is no equivalent reason to serve a user that way, and
# a flag that can be flipped in one place and forgotten in another is how it would happen.
assert agent.GUARDS is True, "refusing to serve with agent.GUARDS = False"

# --- the cache, and where it deliberately is NOT ------------------------------------------
# Phase 6.5's contract is "cache OFF during evals". That is satisfied here STRUCTURALLY, not
# by a flag: the cache is consulted in this file only. run_eval.py and red_team.py call
# agent.run_agent directly and cannot reach it. A flag would have been one forgotten line
# away from the failure lesson 96 records - a harness whose default contradicted the shipped
# configuration for weeks without anyone noticing.
#
# The fingerprint is computed once, at import, and it is part of the cache KEY. A changed
# system therefore cannot reach an old entry at all - the rows become unreachable rather than
# merely stale.
#
# IT USED TO BE THE INDEX ALONE, and that was too narrow in a way nothing would ever have
# reported. `cache.index_fingerprint(agent.FILINGS)` digests the (company, period) pairs and
# nothing else, so editing the answer prompt, changing the model, or changing MAX_ROUNDS left
# it exactly where it was - and every cached answer from the old configuration went on being
# served, silently, because nothing downstream of a cache disagrees with it. Worse, the
# machinery for the narrowest case was already written and never wired: index_fingerprint
# takes a `chunk_count` and this line never passed one, so re-chunking the SAME six filings
# produced the SAME fingerprint. Phase 7, version.py.
#
# The COMPONENTS are kept alongside the digest, not thrown away, so /health can say WHICH part
# moved. "The version changed" sends someone diffing two deployments by hand; "the answer
# prompt moved and the index did not" is an answer.
RUN_COMPONENTS = version.components(
    # The template ACTUALLY SELECTED, which folds in GUARD_PROMPT and SHOW_ARITHMETIC without
    # hashing either separately - and correctly leaves HARDENED_PROMPT out of the digest while
    # it ships disabled, so editing an unused prompt cannot invalidate a single cached answer.
    answer_prompt=agent._with_arithmetic_rule(
        agent.guards.HARDENED_PROMPT if agent.GUARD_PROMPT else rag.PROMPT),
    plan_prompt=agent.PLAN_PROMPT,
    reflect_prompt=agent.REFLECT_PROMPT,
    rewrite_prompt=rewriter.REWRITE_PROMPT,
    # Read off the bound model object, not typed again here - see
    # agent.model_identity() for why a second copy of this fact is a defect.
    model=agent.model_identity()[0],
    temperature=agent.model_identity()[1],
    # Every bound that shapes retrieval. These live in code and never in a prompt (the Phase
    # 4.1 rule), which is also what makes them hashable at all.
    bounds={"MAX_JOBS": agent.MAX_JOBS, "MAX_ROUNDS": agent.MAX_ROUNDS,
            "K_PER_JOB": agent.K_PER_JOB, "PER_JOB_FLOOR": agent.PER_JOB_FLOOR,
            "MIN_CHUNKS": agent.MIN_CHUNKS, "FOLLOWUP_PER_JOB": agent.FOLLOWUP_PER_JOB,
            "MAX_REFLECT_JOBS": agent.MAX_REFLECT_JOBS, "GUARDS": agent.GUARDS,
            "SEARCH_ATTEMPTS": agent.SEARCH_ATTEMPTS},
    filings=agent.FILINGS,
    chunk_count=agent.CHUNK_COUNT,
)
CACHE_FINGERPRINT = version.run_version(RUN_COMPONENTS)
CACHE_ENABLED = os.environ.get("CACHE_OFF", "").lower() not in ("1", "true", "yes")


# --- /ask rate limiting -----------------------------------------------------------------------
# WHY THIS EXISTS, AND WHY IT TOOK THIS LONG. "/ask is not rate limited; login is" has been a
# WRITTEN limitation since 6.2 - a signed-in user can spend the API key as fast as they can
# type. That was acceptable for one analyst on a laptop and is not acceptable anywhere else,
# and it was written down precisely so a later phase could not discover it in a bill.
#
# TWO KNOBS, because they stop two different failures and neither one covers the other:
#   - a tight loop is a BURST, and only a COUNT catches it;
#   - a slow drain is a BUDGET, and only MONEY catches it. Forty cheap questions and four
#     expensive ones cost the same, and a count alone cannot tell them apart.
#
# Both live HERE, in code - not in a prompt, not in a settings table, not in an environment
# variable that can be quietly widened. Same rule as MAX_JOBS and MAX_ROUNDS: a bound that
# lives somewhere it can be argued with is not a bound.
#
# WHAT THIS DOES NOT DO, stated so it is not trusted past its limits: it is PER USER. It does
# nothing about many accounts, and signup is not rate limited by account creation rate. That
# is a different defence and it is not built.
ASK_WINDOW_MINUTES = 60
MAX_PAID_ASKS_PER_WINDOW = 40
MAX_USD_PER_WINDOW = 1.00


# --- database, one connection per request -----------------------------------------------------
# sqlite3 connections are not safe to share across threads, and every endpoint below is a
# plain `def`, which FastAPI runs in a worker thread. One connection per request is the
# boring correct answer; WAL mode (set in db.connect) keeps readers from blocking the writer.

def get_db():
    conn = db.connect()
    try:
        yield conn
    finally:
        conn.close()


with db.connect() as _boot:
    db.init_db(_boot)


# --- auth plumbing ------------------------------------------------------------------------------
#
# WHY THE SESSION COOKIE IS NOT SIGNED, departing from the plan written in the tracker.
# A signature protects a cookie whose VALUE CARRIES MEANING the server would otherwise trust -
# `user_id=7`, `role=admin`. Ours carries no meaning: it is 32 bytes from secrets.token_urlsafe,
# and the server resolves it by looking up its SHA-256 in the sessions table. Tampering with it
# produces a token that fails that lookup, which is exactly what a bad signature would produce,
# one step earlier. Signing it would add a second secret to generate, store and rotate, in
# exchange for nothing. Opaque-random-plus-server-lookup and signed-payload are two designs that
# solve the same problem; using both is ceremony, and ceremony is where real defences go to hide.
#
# HttpOnly is not optional and is not about tampering: it stops page JavaScript from reading the
# token at all, which is what turns an XSS bug from "session stolen" into "session not stolen".

def _set_session_cookies(response: Response, token: str) -> str:
    csrf = secrets.token_urlsafe(24)
    response.set_cookie(SESSION_COOKIE, token, httponly=True, samesite="lax",
                        secure=COOKIE_SECURE, max_age=db.SESSION_DAYS * 86400, path="/")
    # NOT HttpOnly, deliberately: the front end has to read this one to echo it back in a
    # header. That is the whole double-submit mechanism, and it is safe precisely because
    # the value proves nothing on its own - it only has to MATCH.
    response.set_cookie(CSRF_COOKIE, csrf, httponly=False, samesite="lax",
                        secure=COOKIE_SECURE, max_age=db.SESSION_DAYS * 86400, path="/")
    return csrf


def _clear_cookies(response: Response):
    for name in (SESSION_COOKIE, CSRF_COOKIE):
        response.delete_cookie(name, path="/")


def current_user(conn=Depends(get_db), fa_session: Optional[str] = Cookie(default=None)):
    user = auth.current_user(conn, fa_session)
    if user is None:
        raise HTTPException(401, "Not signed in.")
    return user


def csrf_guard(request: Request,
               fa_csrf: Optional[str] = Cookie(default=None),
               x_csrf_token: Optional[str] = Header(default=None)):
    """Double-submit: the cookie and the header must match on every state-changing request.

    An attacker's page can MAKE your browser send a request carrying your cookies, but the
    same-origin policy stops it READING those cookies, so it cannot put the value in a
    header. SameSite=Lax already blocks the classic cross-site form POST; this is the second
    lock, for the cases SameSite does not cover.

    Rejects on absence as well as mismatch. A check that passes when the token is missing is
    not a check - and "missing" is the state an attacker can actually produce.
    """
    if request.method in ("GET", "HEAD", "OPTIONS"):
        return
    if not fa_csrf or not x_csrf_token or not secrets.compare_digest(fa_csrf, x_csrf_token):
        raise HTTPException(403, "CSRF token missing or invalid.")


def _client_ip(request: Request) -> str:
    return request.client.host if request.client else "unknown"


# --- schemas ----------------------------------------------------------------------------------

class SignupIn(BaseModel):
    email: str
    password: str
    display_name: Optional[str] = None


class LoginIn(BaseModel):
    email: str
    password: str


class AskIn(BaseModel):
    question: str = Field(min_length=1, max_length=2000)
    conversation_id: Optional[int] = None


# --- errors -------------------------------------------------------------------------------------
# One handler, so an auth failure cannot accidentally leak through a route that forgot to
# catch it. The message shown is the exception's own text, which for InvalidCredentials is
# identical for "wrong password" and "no such account" by construction (see auth.py).

@app.exception_handler(auth.AuthError)
def _auth_error(request: Request, exc: auth.AuthError):
    status = {"rate_limited": 429, "invalid_credentials": 401,
              "email_taken": 409}.get(exc.code, 400)
    return JSONResponse({"error": exc.code, "detail": str(exc)}, status_code=status)


# --- routes: health and corpus -------------------------------------------------------------------

# --- the interface ---------------------------------------------------------------------------
# One file, read from disk on each request so a UI edit needs no restart. No framework, no
# build step, no CDN: the whole front end is `uvicorn app:app` and a browser, which is what
# makes it runnable by someone who did not write it.
_UI = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ui.html")


@app.get("/")
def index():
    return FileResponse(_UI, media_type="text/html")


@app.get("/health")
def health():
    """Deliberately reports the CONFIGURATION, not just 'ok'.

    Phase 6.0 found a red-team harness that had been silently testing a prompt layer which
    ships disabled, and the only reason it was caught is that the harness printed the config
    it was about to test. A health check that says "ok" tells you the process is alive; this
    one tells you which system is alive.
    """
    return {"status": "ok", "filings": len(agent.FILINGS),
            # Reported for the same reason as everything else here: a limit nobody can read
            # is a limit nobody can check. These three are the whole /ask rate limit.
            "ask_window_minutes": ASK_WINDOW_MINUTES,
            "max_paid_asks_per_window": MAX_PAID_ASKS_PER_WINDOW,
            "max_usd_per_window": MAX_USD_PER_WINDOW,
            "cache_enabled": CACHE_ENABLED, "cache_fingerprint": CACHE_FINGERPRINT,
            # The components, not just the digest. Two deployments that disagree can
            # be compared line by line without anyone diffing source trees by hand,
            # which is the only reason a version number is worth printing at all.
            "run": version.describe(RUN_COMPONENTS),
            "guards": agent.GUARDS, "guard_prompt": agent.GUARD_PROMPT,
            "schema_version": db.SCHEMA_VERSION, "cookie_secure": COOKIE_SECURE,
            # Reported because it can be off for a reason nobody remembers. On 2026-08-18 the
            # LangSmith monthly trace quota ran out and every call started printing a 429 -
            # tracing was already dead at that point, and the only thing the flag still bought
            # was noise. A capability that is off must be VISIBLE, or it comes back as
            # "why are there no traces for last month".
            "langsmith_tracing":
                os.environ.get("LANGSMITH_TRACING", "").lower() in ("1", "true", "yes")}


@app.get("/filings")
def filings():
    """Read from the index at import, never hardcoded - the same list the planner is given.

    A hardcoded corpus description is a second source of truth for a fact the index already
    owns, which is precisely the defect Phase 4.3 removed from PLAN_PROMPT.
    """
    return {"filings": [{"company": c, "period": p} for c, p in agent.FILINGS]}


# --- routes: auth ---------------------------------------------------------------------------------

@app.post("/auth/signup")
def signup(body: SignupIn, response: Response, request: Request, conn=Depends(get_db)):
    uid = auth.signup(conn, body.email, body.password, body.display_name)
    token = db.create_session(conn, uid, user_agent=request.headers.get("user-agent"))
    db.touch_login(conn, uid)
    csrf = _set_session_cookies(response, token)
    user = db.get_user_by_email(conn, body.email)
    return {"user": {"id": user["id"], "email": user["email"],
                     "display_name": user["display_name"]}, "csrf_token": csrf}


@app.post("/auth/login")
def login(body: LoginIn, response: Response, request: Request, conn=Depends(get_db)):
    user, token = auth.login(conn, body.email, body.password,
                             ip=_client_ip(request),
                             user_agent=request.headers.get("user-agent"))
    csrf = _set_session_cookies(response, token)
    return {"user": {"id": user["id"], "email": user["email"],
                     "display_name": user["display_name"]}, "csrf_token": csrf}


@app.post("/auth/logout")
def logout(response: Response, conn=Depends(get_db),
           fa_session: Optional[str] = Cookie(default=None),
           _=Depends(csrf_guard)):
    auth.logout(conn, fa_session)
    _clear_cookies(response)
    return {"ok": True}


@app.get("/auth/me")
def me(user=Depends(current_user), conn=Depends(get_db)):
    return {"user": {"id": user["id"], "email": user["email"],
                     "display_name": user["display_name"]},
            "spend": db.user_spend(conn, user["id"])}


# --- routes: conversations --------------------------------------------------------------------------

def _owned_conversation(conn, conversation_id, user_id):
    """Fetch a conversation, or 404 if it is not this user's.

    404 rather than 403, and the ownership check is here rather than repeated at each call
    site. 403 would confirm that the id exists and belongs to somebody, which is the same
    class of leak as telling a stranger which email addresses have accounts.
    """
    row = conn.execute("SELECT * FROM conversations WHERE id = ? AND user_id = ?",
                       (conversation_id, user_id)).fetchone()
    if row is None:
        raise HTTPException(404, "No such conversation.")
    return row


@app.get("/conversations")
def list_conversations(user=Depends(current_user), conn=Depends(get_db)):
    rows = db.list_conversations(conn, user["id"])
    return {"conversations": [dict(r) for r in rows]}


@app.get("/conversations/{conversation_id}")
def get_conversation(conversation_id: int, user=Depends(current_user), conn=Depends(get_db)):
    conv = _owned_conversation(conn, conversation_id, user["id"])
    msgs = db.list_messages(conn, conversation_id)
    return {"conversation": dict(conv), "messages": [dict(m) for m in msgs]}


@app.delete("/conversations/{conversation_id}")
def archive_conversation(conversation_id: int, user=Depends(current_user),
                         conn=Depends(get_db), _=Depends(csrf_guard)):
    """Archive, not delete. The trace rows attached to this conversation are the only record
    of what the system did and what it cost; a hard delete would destroy measurement data to
    tidy a sidebar."""
    _owned_conversation(conn, conversation_id, user["id"])
    conn.execute("UPDATE conversations SET archived_at = ? WHERE id = ?",
                 (db.utcnow(), conversation_id))
    conn.commit()
    return {"ok": True}


# --- export ---------------------------------------------------------------------------------
# THE TRANSCRIPT ONLY: questions, answers, timestamps. No sources, no pipeline, no tokens, no
# cost. That is what was asked for and it is also the right split - the trace panel is an
# engineering artefact, and a document mixing it into the conversation would be neither a
# clean record nor a usable audit trail. If an audit export is ever wanted it should be its
# own endpoint with its own format, not this one with a flag.
#
# Both formats are generated in memory and streamed; nothing is written to disk, so there is
# no temp file to leak between users on a shared machine.

def _transcript(conn, conversation_id, user_id):
    conv = _owned_conversation(conn, conversation_id, user_id)
    rows = db.list_messages(conn, conversation_id)
    return conv, [(r["role"], r["content"], r["created_at"]) for r in rows]


def _safe_name(title, ext):
    keep = "".join(ch if (ch.isalnum() or ch in " -_") else "" for ch in (title or "chat"))
    return (keep.strip()[:60] or "conversation") + ext


@app.get("/conversations/{conversation_id}/export")
def export_conversation(conversation_id: int, format: str = "docx",
                        user=Depends(current_user), conn=Depends(get_db)):
    import io
    conv, msgs = _transcript(conn, conversation_id, user["id"])
    fmt = (format or "docx").lower()
    if fmt not in ("docx", "pdf"):
        raise HTTPException(400, "format must be docx or pdf")
    buf = io.BytesIO()

    if fmt == "docx":
        try:
            from docx import Document as Docx
            from docx.shared import Pt, RGBColor
        except ImportError:
            raise HTTPException(
                501, "python-docx is not installed. Run: pip install python-docx")
        doc = Docx()
        doc.add_heading(conv["title"], level=1)
        p = doc.add_paragraph(f"Financial Research & Compliance Analyst · exported "
                              f"{db.utcnow()}")
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x98, 0xA2, 0xB3)
        for role, content, created in msgs:
            head = doc.add_paragraph()
            r = head.add_run(("You" if role == "user" else "Analyst") + "  ·  " + created)
            r.bold = True
            r.font.size = Pt(9)
            body = doc.add_paragraph(content)
            body.paragraph_format.space_after = Pt(12)
        doc.save(buf)
        media = ("application/vnd.openxmlformats-officedocument"
                 ".wordprocessingml.document")
    else:
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError:
            raise HTTPException(501, "reportlab is not installed. Run: pip install reportlab")
        ss = getSampleStyleSheet()
        who = ParagraphStyle("who", parent=ss["Normal"], fontSize=8, textColor="#98A2B3",
                             spaceAfter=2)
        body = ParagraphStyle("body", parent=ss["Normal"], fontSize=10.5, leading=15,
                              spaceAfter=12)
        doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=18 * mm, bottomMargin=18 * mm,
                                leftMargin=20 * mm, rightMargin=20 * mm,
                                title=conv["title"])
        flow = [Paragraph(_x(conv["title"]), ss["Heading1"]),
                Paragraph("Financial Research &amp; Compliance Analyst · exported "
                          + db.utcnow(), who),
                Spacer(1, 8)]
        for role, content, created in msgs:
            flow.append(Paragraph(("You" if role == "user" else "Analyst")
                                  + " &nbsp;·&nbsp; " + created, who))
            flow.append(Paragraph(_x(content).replace("\n", "<br/>"), body))
        doc.build(flow)
        media = "application/pdf"

    buf.seek(0)
    name = _safe_name(conv["title"], "." + fmt)
    return StreamingResponse(buf, media_type=media, headers={
        "Content-Disposition": f'attachment; filename="{name}"'})


@app.get("/messages/{message_id}/trace")
def get_trace(message_id: int, user=Depends(current_user), conn=Depends(get_db)):
    """Everything the trace panel shows, exactly as stored. No arithmetic happens here.

    That is the contract from 6.1: the panel is a SELECT. If this endpoint recomputed a
    total, the product and the eval could disagree while both looked self-consistent.
    """
    trace = db.get_trace(conn, message_id)
    if trace is None or trace["user_id"] != user["id"]:
        raise HTTPException(404, "No such trace.")
    return trace


# --- admission to the paid path -------------------------------------------------------------------
#
# WHAT WAS MEASURED, and it is worse than the defect that was predicted. probe_concurrency.py
# lowered the limit to leave a budget of FOUR paid questions and then sent fourteen requests at
# the same instant. Fourteen were answered. **Zero were refused.** All fourteen were charged.
#
# The reason is not subtle once it is written down. The old guard read the count, then the
# agent ran, and only THEN was the row that records the charge written. Every request that
# arrived inside that gap read a count that did not yet include any of the others - so they
# all saw the same number and they all passed. The gap is not a scheduling accident measured
# in microseconds: it is the whole agent call, roughly eight seconds in production. A limit
# that is CHECKED and then ACTED ON, with the evidence of the act written last, is not a
# limit. It is a suggestion with good intentions.
#
# THE FIX IS TO MAKE THE CLAIM ITSELF THE FIRST WRITE. A request reserves its slot under a
# lock, before any model call, and the reservation is visible to every other request from that
# instant. The count that admission compares against is therefore
#
#     paid answers already RECORDED   +   paid answers currently IN FLIGHT
#
# and the second term is the one the old code could not see. The lock is held only for the
# arithmetic - never across the agent call - so this serialises admission, not answering.
#
# WHAT THIS IS NOT: process-wide, and that limitation is real rather than hidden. Two uvicorn
# workers would keep two separate counters and the bound would be per-worker. That is exactly
# the constraint the Dockerfile already pins with `--workers 1`, for the SQLite reason, so the
# validity of this defence and the shape of the deployment are the same shape. If this ever
# runs multi-process the reservation has to move into the database, and the honest place to
# discover that is here rather than in production.
#
# AND THE SPEND BOUND STAYS APPROXIMATE, deliberately. A request's cost is not known until it
# finishes, so there is nothing truthful to reserve; MAX_USD_PER_WINDOW is still compared
# against what has been recorded. The COUNT bound is now exact and the SPEND bound is a
# trailing one, and saying so is better than implying both were fixed.

_ASK_INFLIGHT_LOCK = threading.Lock()
_ASK_INFLIGHT = {}          # user_id -> paid asks admitted but not yet recorded


def ask_slot(user=Depends(current_user), conn=Depends(get_db)):
    """Admit one paid ask, or refuse it. Releases the slot however the request ends.

    A dependency with `yield` rather than a block inside the endpoint, because the release
    MUST happen on every exit - a 500 from the agent, a guard that raises, a client that
    disconnects. A `finally` that lives in the same function as two hundred lines of answering
    is a `finally` somebody eventually returns past; FastAPI runs this teardown for us.
    """
    uid = user["id"]
    with _ASK_INFLIGHT_LOCK:
        used = db.recent_asks(conn, uid, ASK_WINDOW_MINUTES)
        pending = _ASK_INFLIGHT.get(uid, 0)
        if used["paid"] + pending >= MAX_PAID_ASKS_PER_WINDOW:
            # THE WORDING IS PRECISE ON PURPOSE, and the first draft of it was wrong. A cache
            # hit does not COUNT towards the limit, but once the limit is reached it is not
            # EXEMPT from it either - this guard sits ahead of the cache lookup, because the
            # lookup needs the rewriter to have run and the rewriter is itself a paid call.
            # Saying "cached repeats are free" here would have been a false statement printed
            # at the exact moment a user is trying to understand why they were refused.
            raise HTTPException(429, f"Rate limit reached: {MAX_PAID_ASKS_PER_WINDOW} paid "
                                     f"questions per {ASK_WINDOW_MINUTES} minutes. Answers "
                                     f"served from the cache are not counted towards it, but "
                                     f"no further requests are accepted until the window "
                                     f"clears.")
        if used["usd"] >= MAX_USD_PER_WINDOW:
            raise HTTPException(429, f"Spend limit reached: ${MAX_USD_PER_WINDOW:.2f} per "
                                     f"{ASK_WINDOW_MINUTES} minutes. Try again shortly.")
        _ASK_INFLIGHT[uid] = pending + 1
    try:
        yield
    finally:
        with _ASK_INFLIGHT_LOCK:
            remaining = _ASK_INFLIGHT.get(uid, 1) - 1
            if remaining > 0:
                _ASK_INFLIGHT[uid] = remaining
            else:
                # POPPED, not set to zero. A dict that keeps a 0 for every account that has
                # ever asked a question is a slow leak whose only symptom is memory, and this
                # process is meant to stay up.
                _ASK_INFLIGHT.pop(uid, None)


# --- the one that matters ---------------------------------------------------------------------------

@app.post("/ask")
def ask(body: AskIn, request: Request, user=Depends(current_user), conn=Depends(get_db),
        _=Depends(csrf_guard), _slot=Depends(ask_slot)):
    """Answer one question, store it, and store how it was answered.

    A plain `def`, not `async def`, and that is load-bearing. agent.run_agent blocks for
    roughly eight seconds of network I/O in synchronous libraries. Declared `async def` it
    would run ON the event loop and every other request in the process - including the UI
    polling for history - would freeze until it finished. FastAPI runs sync endpoints in a
    worker thread pool, which is the correct place for blocking work.
    """
    question = body.question.strip()
    if not question:
        raise HTTPException(400, "Question is empty.")

    # The rate limit is no longer checked here. It is ADMISSION, and it happens in the
    # ask_slot dependency above, before this function is entered. See that function for the
    # measurement that moved it.

    if body.conversation_id is None:
        title = question[:60] + ("…" if len(question) > 60 else "")
        conversation_id = db.create_conversation(conn, user["id"], title)
        history = []
    else:
        conversation_id = _owned_conversation(conn, body.conversation_id, user["id"])["id"]
        # BEFORE the new message is stored, or the question ends up inside its own history.
        history = db.history_pairs(conn, conversation_id)

    db.add_message(conn, conversation_id, "user", question)

    # --- rewrite, then cache, then answer ---------------------------------------------------
    # THE ORDER IS THE POINT, and getting it wrong was a live defect found by using the UI
    # rather than by any test:
    #
    #   1. This endpoint used to call run_agent WITHOUT history, so the 6.4 rewriter never ran
    #      in the product at all. "And what was its net income?" reached the pipeline verbatim
    #      and came back with "please specify which company". A feature built, measured at
    #      14/14 in its own eval, and never wired to anything.
    #   2. The cache was keyed on the RAW question, so "And AMD?" became a global cache key -
    #      the exact catastrophic key the 6.5 plan named in advance, shipped by accident. The
    #      comment here even claimed the opposite of what the code did, which is worse than
    #      having no comment.
    #
    # So the rewrite happens HERE, in front of the cache, rather than inside run_agent: the
    # cache sits before the agent, and it must be keyed on the standalone question. run_agent
    # is then called with no history, so nothing is rewritten twice.
    import time
    t0 = time.perf_counter()
    with telemetry.capture() as calls:
        asked, rewrite_note = rewriter.rewrite(question, history,
                                               filings=agent.FILINGS)
        rewritten = asked if history else None

        hit = cache.get(conn, asked, CACHE_FINGERPRINT) if CACHE_ENABLED else None
        if hit is None:
            out = agent.run_agent(question=asked)

    if hit is not None:
        message_id = db.add_message(conn, conversation_id, "assistant", hit["answer"])
        db.save_trace(
            conn, message_id=message_id, conversation_id=conversation_id, user_id=user["id"],
        # WHICH SYSTEM ANSWERED. Recorded per row rather than looked up later, because
        # "later" is after the configuration has already changed - which is the only
        # time anybody asks.
        run_version=CACHE_FINGERPRINT,
            question_raw=question, question_rewritten=rewritten,
            jobs=json.loads(hit["jobs_json"]) if hit["jobs_json"] else None,
            filings=json.loads(hit["filings_json"]) if hit["filings_json"] else None,
            chunk_ids=json.loads(hit["chunk_ids_json"]) if hit["chunk_ids_json"] else None,
            retrieval=json.loads(hit["retrieval_json"]) if hit["retrieval_json"] else None,
            rounds=hit["rounds"], reflect_fired=False, guard_fired="", cache_hit=True,
            # A cache hit on a SINGLE-TURN question is free. A hit on a follow-up is not: the
            # rewrite was paid for before the lookup could happen, and recording it as zero
            # would understate what this product costs.
            calls=telemetry.rows_for_db(calls),
            # ...and what it would have cost WITHOUT the cache, read off the stored row rather
            # than estimated. This is the only honest way for the panel to claim a saving:
            # the number is what the first run of this exact question actually paid.
            saved_usd=hit["usd"], saved_input_tokens=hit["input_tokens"],
            saved_output_tokens=hit["output_tokens"],
            seconds=round(time.perf_counter() - t0, 3))
        return {"conversation_id": conversation_id, "message_id": message_id,
                "answer": hit["answer"], "context": hit["context"],
                "rounds": hit["rounds"], "guard_fired": "", "cache_hit": True,
                "cached_is_refusal": bool(hit["is_refusal"]),
                "cached_at": hit["created_at"],
                "rewrite_note": rewrite_note,
                "trace": db.get_trace(conn, message_id)}

    # THE CONTRACT. run_agent above is the exact function run_eval.py calls, inside the exact
    # capture run_eval.py uses.
    seconds = round(time.perf_counter() - t0, 3)

    answer = out["answer"]
    message_id = db.add_message(conn, conversation_id, "assistant", answer)

    chunks = out.get("chunks") or []
    chunk_ids = [getattr(d, "id", None) or (getattr(d, "metadata", {}) or {}).get("id")
                 for d in chunks]
    filings_used = sorted({
        f"{(d.metadata or {}).get('company', '?')} | {(d.metadata or {}).get('period', '?')}"
        for d in chunks if getattr(d, "metadata", None)})

    db.save_trace(
        conn, message_id=message_id, conversation_id=conversation_id, user_id=user["id"],
        # WHICH SYSTEM ANSWERED. Recorded per row rather than looked up later, because
        # "later" is after the configuration has already changed - which is the only
        # time anybody asks.
        run_version=CACHE_FINGERPRINT,
        question_raw=question,
        question_rewritten=rewritten,
        # jobs_log, NOT jobs. `jobs` is a work queue that retrieve_node drains, so by the
        # time the graph returns it is empty and the panel would show "no plan" on every
        # single question - a blank that looks like a UI bug and is actually a lost record.
        jobs=out.get("jobs_log") or None,
        filings=filings_used,
        chunk_ids=[c for c in chunk_ids if c],
        rounds=out.get("rounds"),
        # Reflect fired iff a second retrieval round happened. Recorded as an observation of
        # the run, not inferred in the UI later - see the 6.1 contract.
        reflect_fired=bool(out.get("rounds") and out["rounds"] > 1),
        guard_fired=out.get("guard_fired") or "",
        retrieval=out.get("retrieval_log") or None,
        cache_hit=False,
        calls=telemetry.rows_for_db(calls),
        seconds=seconds)
    totals = telemetry.summarise(calls)

    # Store only after the guards have had their say. cacheable() refuses anything produced
    # under attack or any refusal, and it says which - recorded rather than silent, so a cache
    # that is never filling up is visible instead of merely disappointing.
    stored, why_not = (cache.put(
        # `asked`, not `question`: the key must be the standalone question, never the raw
        # follow-up. See the ordering note above.
        conn, question=asked, fingerprint=CACHE_FINGERPRINT, answer=answer,
        context=out["context"], guard_fired=out.get("guard_fired") or "",
        jobs=out.get("jobs_log") or None, filings=filings_used,
        chunk_ids=[c for c in chunk_ids if c], rounds=out.get("rounds"),
        retrieval=out.get("retrieval_log") or None,
        # 6.7: a refusal IS stored now, but not one from a run whose planner fell back.
        degraded=out.get("degraded") or "",
        input_tokens=totals["input_tokens"], output_tokens=totals["output_tokens"],
        usd=totals["usd"], user_id=user["id"]) if CACHE_ENABLED else (False, "cache disabled"))

    return {
        "conversation_id": conversation_id,
        "message_id": message_id,
        "cache_hit": False,
        "cached": stored, "not_cached_because": why_not,
        "rewrite_note": rewrite_note,
        # the same fields run_eval.py scores, under the same names the agent returns them by
        "answer": answer,
        "context": out["context"],
        "rounds": out.get("rounds"),
        "guard_fired": out.get("guard_fired") or "",
        "trace": db.get_trace(conn, message_id),
    }


if __name__ == "__main__":
    import uvicorn

    # The app OBJECT, not the string "app:app". uvicorn resolves an import string by importing
    # the module - and this module is already running, under the name `__main__`. Python's
    # module cache has no entry for `app`, so it would execute this whole file a SECOND time:
    # two FastAPI instances, two init_db calls, two of every module-level side effect, and the
    # one that is served is not the one whose globals `__main__` just set up.
    #
    # The import-string form is only worth its cost when reload=True, which needs a name it can
    # re-import on every file change. Serving is reload=False, so the object is strictly better.
    uvicorn.run(app, host="127.0.0.1", port=8000)
